import argparse
import logging
import multiprocessing
import re
import shutil
import subprocess
import sys
import tempfile
import unicodedata
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path

try:
    from pdfminer.high_level import extract_text as pdfminer_extract
except ImportError:  # pragma: no cover - handled at runtime
    pdfminer_extract = None

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:  # pragma: no cover - handled at runtime
    Document = None

# ---------------------------------------------------------------------------
# Logging
# ---------------------------------------------------------------------------

LOG = logging.getLogger("nepali-pdf")

REQUIRED_COMMANDS = ("pdffonts", "pdftotext", "ocrmypdf")

# Patterns that indicate garbled / mis-extracted Devanagari text.
BAD_PATTERNS = [
    r"ष्ट्र",
    r"प्",
    r"अतभ",
    r"त्तर्",
    r"ाँ",
    # Lone combining marks (not preceded by a base letter)
    r"(?<!\u0915-\u0939])[\u093E-\u094D]",
    # Repeated virama clusters (◌्◌्)
    r"\u094D{2,}",
    # Common garbled sequences from bad font mappings
    r"ि◌",
    r"◌ा",
]

# Unicode range for Devanagari block: U+0900 – U+097F
DEVANAGARI_RANGE = re.compile(r"[\u0900-\u097F]")

# Characters considered "weird" punctuation / control that shouldn't dominate.
_JUNK_CHARS = re.compile(r"[^\w\s\u0900-\u097F.,;:!?\-()\"'०-९]", re.UNICODE)

# Subprocess timeout in seconds – prevents a single hung PDF from blocking
# the entire batch forever.
SUBPROCESS_TIMEOUT = 300


def ensure_python_dependencies() -> None:
    if pdfminer_extract is None:
        raise RuntimeError(
            "Missing Python dependency 'pdfminer.six'. Install it with "
            "'pip install -r requirements.txt'."
        )
    if Document is None:
        raise RuntimeError(
            "Missing Python dependency 'python-docx'. Install it with "
            "'pip install -r requirements.txt'."
        )


def ensure_external_dependencies() -> None:
    missing_commands = [command for command in REQUIRED_COMMANDS if shutil.which(command) is None]
    if missing_commands:
        joined = ", ".join(missing_commands)
        raise RuntimeError(
            f"Missing required system commands: {joined}. "
            "Install poppler-utils and ocrmypdf (plus Tesseract language packs)."
        )


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def check_unicode_mapping(pdf_path: Path) -> bool:
    """Return True when font metadata suggests Unicode text extraction is possible."""
    try:
        result = subprocess.run(
            ["pdffonts", str(pdf_path)],
            capture_output=True,
            text=True,
            check=True,
            timeout=SUBPROCESS_TIMEOUT,
        )
    except subprocess.TimeoutExpired:
        LOG.warning("pdffonts timed out for %s – assuming no Unicode mapping.", pdf_path.name)
        return False
    except subprocess.CalledProcessError as exc:
        details = exc.stderr.strip() or exc.stdout.strip() or str(exc)
        raise RuntimeError(f"Font inspection failed: {details}") from exc

    lines = result.stdout.splitlines()
    for line in lines[2:]:
        if not line.strip():
            continue
        columns = line.split()
        if len(columns) >= 6 and columns[5].lower() == "no":
            return False
    return True


def extract_with_poppler(pdf_path: Path) -> str:
    """Extract text with pdftotext and preserve layout where possible."""
    try:
        result = subprocess.run(
            ["pdftotext", "-layout", "-enc", "UTF-8", str(pdf_path), "-"],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            check=True,
            timeout=SUBPROCESS_TIMEOUT,
        )
    except subprocess.TimeoutExpired:
        raise RuntimeError(f"pdftotext timed out after {SUBPROCESS_TIMEOUT}s.")
    except subprocess.CalledProcessError as exc:
        details = exc.stderr.strip() or exc.stdout.strip() or str(exc)
        raise RuntimeError(f"Poppler extraction failed: {details}") from exc

    return result.stdout


# ---------------------------------------------------------------------------
# Text quality scoring
# ---------------------------------------------------------------------------

# Precalculate a translation table that maps all whitespace to None
_WS_REMOVAL = str.maketrans({c: None for c in " \t\n\r\f\v"})


def _devanagari_ratio(text: str) -> float:
    """Return the fraction of non-whitespace characters that are Devanagari."""
    non_ws = text.translate(_WS_REMOVAL)
    if not non_ws:
        return 0.0
    devanagari_count = len(DEVANAGARI_RANGE.findall(non_ws))
    return devanagari_count / len(non_ws)


def _ascii_alnum_ratio(text: str) -> float:
    """Return the fraction of non-whitespace characters that are ASCII A-Za-z0-9."""
    non_ws = text.translate(_WS_REMOVAL)
    if not non_ws:
        return 0.0
    alnum_count = sum(1 for char in non_ws if char.isascii() and char.isalnum())
    return alnum_count / len(non_ws)


def _ascii_printable_ratio(text: str) -> float:
    """Return the fraction of non-whitespace characters that are printable ASCII."""
    non_ws = text.translate(_WS_REMOVAL)
    if not non_ws:
        return 0.0
    printable_count = sum(1 for char in non_ws if 32 <= ord(char) <= 126)
    return printable_count / len(non_ws)


def _junk_ratio(text: str) -> float:
    """Return the fraction of non-whitespace characters that are 'junk'."""
    non_ws = text.translate(_WS_REMOVAL)
    if not non_ws:
        return 0.0
    junk_count = len(_JUNK_CHARS.findall(non_ws))
    return junk_count / len(non_ws)


def is_text_valid(text: str) -> bool:
    """
    Evaluate whether extracted text looks usable.

    Uses a combined scoring approach:
    1. Text must not be empty.
    2. Accept direct-extracted ASCII text layers as-is when the copied text is
       clean English-style characters (`A-Za-z0-9`) even if it represents a
       legacy font encoding such as Preeti.
    3. Only apply the Devanagari Unicode quality checks when the extracted text
       actually contains Devanagari characters.
    4. If > 40 % of characters are junk/control characters, reject.
    5. Bad-pattern count is weighed against Devanagari ratio —
       high Devanagari ratio tolerates more bad patterns.
    """
    stripped = text.strip()
    if not stripped:
        return False

    if "\uFFFD" not in stripped and _ascii_printable_ratio(stripped) >= 0.85:
        return _ascii_alnum_ratio(stripped) >= 0.20

    ratio = _devanagari_ratio(stripped)
    if ratio == 0.0:
        return False

    if ratio < 0.10:
        return False

    if _junk_ratio(stripped) > 0.40:
        return False

    error_count = sum(len(re.findall(pattern, stripped)) for pattern in BAD_PATTERNS)

    # Allow more bad hits when the Devanagari ratio is high.
    allowed_errors = 3 if ratio < 0.5 else 8
    return error_count <= allowed_errors


def post_process_text(text: str) -> str:
    """Clean up extracted text for better readability."""
    # Remove soft hyphens entirely to prevent broken words
    text = text.replace("\u00AD", "")

    # Normalize Unicode characters (NFC form is standard for Devanagari).
    text = unicodedata.normalize("NFC", text)

    # Strip trailing whitespace from each line.
    lines = [line.rstrip() for line in text.splitlines()]

    # Collapse runs of 3+ blank lines into 2.
    cleaned: list[str] = []
    blank_run = 0
    for line in lines:
        if line == "":
            blank_run += 1
            if blank_run <= 2:
                cleaned.append(line)
        else:
            blank_run = 0
            cleaned.append(line)

    return "\n".join(cleaned).strip() + "\n"


# ---------------------------------------------------------------------------
# OCR fallback
# ---------------------------------------------------------------------------

def run_ocr_and_extract(pdf_path: Path) -> str:
    """Run OCRmyPDF to rebuild the text layer, then re-extract via Poppler."""
    LOG.info("Running OCR fallback...")
    with tempfile.TemporaryDirectory() as temp_dir:
        ocr_output = Path(temp_dir) / "ocr_output.pdf"
        base_command = [
            "ocrmypdf",
            "-l",
            "nep+hin+eng",
            "--deskew",
            "--optimize",
            "1",
            "--force-ocr",
            "--image-dpi",
            "300",
            str(pdf_path),
            str(ocr_output),
        ]

        def run_ocr(command: list[str]) -> subprocess.CompletedProcess[str]:
            return subprocess.run(
                command,
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace",
                check=True,
                timeout=SUBPROCESS_TIMEOUT,
            )

        try:
            completed = run_ocr(base_command[:4] + ["--clean-final"] + base_command[4:])
        except subprocess.TimeoutExpired:
            raise RuntimeError(f"OCR timed out after {SUBPROCESS_TIMEOUT}s.")
        except subprocess.CalledProcessError as exc:
            details = exc.stderr.strip() or exc.stdout.strip() or str(exc)
            if "Could not find program 'unpaper'" in details:
                LOG.warning("unpaper not found; retrying OCR without --clean-final.")
                try:
                    completed = run_ocr(base_command)
                except subprocess.TimeoutExpired:
                    raise RuntimeError(f"OCR timed out after {SUBPROCESS_TIMEOUT}s.")
                except subprocess.CalledProcessError as retry_exc:
                    retry_details = (
                        retry_exc.stderr.strip() or retry_exc.stdout.strip() or str(retry_exc)
                    )
                    raise RuntimeError(f"OCR fallback failed: {retry_details}") from retry_exc
            else:
                raise RuntimeError(f"OCR fallback failed: {details}") from exc

        if completed.stderr.strip():
            LOG.debug(completed.stderr.strip())
        return extract_with_poppler(ocr_output)


def _pdfminer_worker(pdf_path: str, return_dict: dict) -> None:
    """Target function for the PDFMiner process."""
    try:
        if pdfminer_extract is not None:
            return_dict["text"] = pdfminer_extract(pdf_path)
    except Exception as exc:
        return_dict["error"] = exc


def _run_pdfminer_with_timeout(pdf_path: Path, timeout: int = SUBPROCESS_TIMEOUT) -> str:
    """Run pdfminer in a separate process to allow strict timeouts and avoid hangs."""
    manager = multiprocessing.Manager()
    return_dict = manager.dict()
    p = multiprocessing.Process(target=_pdfminer_worker, args=(str(pdf_path), return_dict))
    
    p.start()
    p.join(timeout)
    
    if p.is_alive():
        p.terminate()
        p.join()
        raise RuntimeError(f"PDFMiner timed out after {timeout}s.")
        
    if "error" in return_dict:
        raise return_dict["error"]
        
    return return_dict.get("text", "")


# ---------------------------------------------------------------------------
# Main extraction pipeline
# ---------------------------------------------------------------------------

def process_nepali_pdf(pdf_path: Path) -> tuple[str, str]:
    """Extract text from a single PDF and return (text, method)."""
    LOG.info("Processing: %s", pdf_path)

    if not pdf_path.exists():
        raise FileNotFoundError(f"Cannot find input PDF: {pdf_path}")
    if pdf_path.suffix.lower() != ".pdf":
        raise ValueError(f"Input is not a PDF: {pdf_path}")

    # Skip zero-byte / corrupt files immediately.
    if pdf_path.stat().st_size == 0:
        raise RuntimeError("File is empty (0 bytes).")

    has_unicode = check_unicode_mapping(pdf_path)
    if has_unicode:
        LOG.info("Unicode mapping found. Trying direct extraction first.")
    else:
        LOG.info("Missing Unicode mapping detected. Trying direct extraction before OCR.")

    text = extract_with_poppler(pdf_path)
    if is_text_valid(text):
        return post_process_text(text), "poppler"

    if has_unicode:
        LOG.info("Poppler output failed validation. Trying PDFMiner.")
    elif text.strip():
        LOG.info("Poppler returned copyable text but it was not valid ASCII/Devanagari. Trying PDFMiner.")
    else:
        LOG.info("Poppler returned no usable text. Falling back to OCR.")

    if has_unicode or text.strip():
        try:
            text = _run_pdfminer_with_timeout(pdf_path)
        except Exception as exc:
            LOG.warning("PDFMiner failed or timed out: %s – skipping to OCR.", exc)
            text = ""

        if is_text_valid(text):
            return post_process_text(text), "pdfminer"

        LOG.info("PDFMiner output failed validation. Falling back to OCR.")

    text = run_ocr_and_extract(pdf_path)
    if text.strip():
        return post_process_text(text), "ocr"

    raise RuntimeError("All extraction methods failed or returned empty text.")


# ---------------------------------------------------------------------------
# DOCX output
# ---------------------------------------------------------------------------

def default_output_path(pdf_path: Path) -> Path:
    return pdf_path.with_suffix(".docx")


def write_docx_output(output_path: Path, text: str, source_name: str = "") -> None:
    """Write extracted text to a DOCX file with Devanagari-friendly formatting.

    Consecutive blank lines are collapsed into paragraph spacing rather than
    creating thousands of empty paragraph objects— this significantly reduces
    file size and generation time for large documents.
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # -- Set default style font to Mangal (Devanagari) at 11pt --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Mangal"
    font.size = Pt(11)

    # Add a title derived from the source PDF name.
    if source_name:
        doc.add_heading(source_name, level=1)

    # Split text into paragraphs: merge consecutive blank lines into spacing.
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.strip():
            para = doc.add_paragraph(line)
            run = para.runs[0] if para.runs else para.add_run("")
            run.font.name = "Mangal"
            run.font.size = Pt(11)
        else:
            # Count consecutive blank lines.
            blank_count = 0
            while i < len(lines) and not lines[i].strip():
                blank_count += 1
                i += 1
            # Add a single empty paragraph with extra spacing.
            para = doc.add_paragraph("")
            para_format = para.paragraph_format
            para_format.space_after = Pt(6 * min(blank_count, 3))
            i -= 1  # compensate for the outer i += 1
        i += 1

    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# Single-file mode
# ---------------------------------------------------------------------------

def process_single_file(input_pdf: Path, output_path: Path | None) -> int:
    target_output = output_path or default_output_path(input_pdf)
    text, method = process_nepali_pdf(input_pdf)
    
    if not text.strip():
        LOG.error("Extracted text is empty. Aborting DOCX creation.")
        return 1

    write_docx_output(target_output, text, source_name=input_pdf.stem)
    LOG.info("Extraction successful via %s.", method)
    LOG.info("Saved DOCX to: %s", target_output)
    return 0


# ---------------------------------------------------------------------------
# Batch / parallel mode  (ThreadPoolExecutor — I/O-bound subprocess work)
# ---------------------------------------------------------------------------

def _process_one_pdf(pdf_path: Path, output_path: Path) -> dict:
    """Worker function: process a single PDF and return a result dict."""
    try:
        text, method = process_nepali_pdf(pdf_path)
        if not text.strip():
            raise RuntimeError("Extracted text is empty.")
        write_docx_output(output_path, text, source_name=pdf_path.stem)
        return {"file": pdf_path.name, "output": str(output_path), "method": method, "ok": True}
    except Exception as exc:
        return {"file": pdf_path.name, "ok": False, "error": str(exc)}


def process_directory(input_dir: Path, output_dir: Path | None, workers: int = 4) -> int:
    if not input_dir.exists():
        raise FileNotFoundError(f"Input directory does not exist: {input_dir}")
    if not input_dir.is_dir():
        raise NotADirectoryError(f"Input path is not a directory: {input_dir}")

    pdf_files = sorted(path for path in input_dir.iterdir() if path.is_file() and path.suffix.lower() == ".pdf")
    if not pdf_files:
        raise FileNotFoundError(f"No PDF files found in directory: {input_dir}")

    total = len(pdf_files)
    LOG.info("Found %d PDF(s). Processing with %d worker(s)...\n", total, workers)

    # Build (pdf, output) pairs.
    tasks: list[tuple[Path, Path]] = []
    for pdf_path in pdf_files:
        if output_dir is None:
            target_output = default_output_path(pdf_path)
        else:
            target_output = output_dir / f"{pdf_path.stem}.docx"
        tasks.append((pdf_path, target_output))

    succeeded = 0
    failed = 0
    results_table: list[dict] = []

    with ThreadPoolExecutor(max_workers=workers) as executor:
        future_to_index = {
            executor.submit(_process_one_pdf, pdf, out): idx
            for idx, (pdf, out) in enumerate(tasks)
        }

        for future in as_completed(future_to_index):
            result = future.result()
            results_table.append(result)
            progress = succeeded + failed + 1

            if result["ok"]:
                succeeded += 1
                LOG.info("[%d/%d] ✓ %s → %s (via %s)", progress, total, result["file"], result["output"], result["method"])
            else:
                failed += 1
                LOG.error("[%d/%d] ✗ %s: %s", progress, total, result["file"], result["error"])

    # -- Structured summary table --
    LOG.info("")
    LOG.info("--- Batch Summary ---")
    LOG.info("%-40s  %-8s  %s", "File", "Status", "Detail")
    LOG.info("%-40s  %-8s  %s", "-" * 40, "-" * 8, "-" * 30)
    for r in results_table:
        status = "OK" if r["ok"] else "FAILED"
        detail = r.get("method", r.get("error", ""))
        LOG.info("%-40s  %-8s  %s", r["file"][:40], status, detail)
    LOG.info("")
    LOG.info("Total: %d  |  Succeeded: %d  |  Failed: %d", total, succeeded, failed)
    return 1 if failed else 0


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Extract Nepali text from PDFs using Poppler, PDFMiner, and OCR fallback. Outputs DOCX files."
    )
    parser.add_argument(
        "input_pdf",
        nargs="?",
        help="Path to a single PDF file to process.",
    )
    parser.add_argument(
        "--output",
        help="Optional DOCX output path for single-file mode.",
    )
    parser.add_argument(
        "--input-dir",
        help="Process all PDFs in this directory (non-recursive).",
    )
    parser.add_argument(
        "--output-dir",
        help="Optional directory to receive .docx outputs in batch mode.",
    )
    parser.add_argument(
        "--workers",
        type=int,
        default=4,
        help="Number of parallel workers for batch mode (default: 4).",
    )
    parser.add_argument(
        "--verbose", "-v",
        action="store_true",
        help="Enable debug-level logging output.",
    )
    parser.add_argument(
        "--quiet", "-q",
        action="store_true",
        help="Suppress all output except errors.",
    )
    return parser


def validate_args(args: argparse.Namespace, parser: argparse.ArgumentParser) -> None:
    if bool(args.input_pdf) == bool(args.input_dir):
        parser.error("Specify exactly one of: input_pdf or --input-dir.")
    if args.output and not args.input_pdf:
        parser.error("--output can only be used with single-file mode.")
    if args.output_dir and not args.input_dir:
        parser.error("--output-dir can only be used with --input-dir.")
    if args.workers < 1:
        parser.error("--workers must be at least 1.")


def _configure_logging(verbose: bool = False, quiet: bool = False) -> None:
    """Set up the root logger based on CLI flags."""
    if quiet:
        level = logging.ERROR
    elif verbose:
        level = logging.DEBUG
    else:
        level = logging.INFO

    logging.basicConfig(
        level=level,
        format="%(message)s",
        stream=sys.stderr,
    )


def main() -> int:
    parser = build_parser()
    args = parser.parse_args()
    validate_args(args, parser)
    _configure_logging(verbose=args.verbose, quiet=args.quiet)

    try:
        ensure_python_dependencies()
        ensure_external_dependencies()

        if args.input_pdf:
            output_path = Path(args.output) if args.output else None
            return process_single_file(Path(args.input_pdf), output_path)

        output_dir = Path(args.output_dir) if args.output_dir else None
        return process_directory(Path(args.input_dir), output_dir, workers=args.workers)
    except Exception as exc:
        LOG.error("Error: %s", exc)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
